#!/usr/bin/env python3
"""Generate DMC / DMP / SAP sample templates and register them in SQLite."""
from __future__ import annotations

import json
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from docx import Document

from config import settings
from db import create_template, get_conn, init_db, list_templates, update_template_mappings
from models.field_schema import FIELD_DEFINITIONS
from services.docx_parser import parse_docx

SEED_DIR = settings.data_dir / "seed_templates"
VALID_KEYS = {f["key"] for f in FIELD_DEFINITIONS}


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _add_line(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def _build_dmc() -> Path:
    path = SEED_DIR / "crogo_dmc_charter.docx"
    doc = Document()
    _add_heading(doc, "数据监查委员会章程 (DMC Charter)", 0)
    _add_line(doc, "方案编号：{{protocol_id}}")
    _add_line(doc, "研究标题：{{study_title}}")
    _add_line(doc, "申办方：{{sponsor}}")
    _add_line(doc, "适应症：{{indication}}")
    _add_heading(doc, "1. 研究概述", 1)
    _add_line(doc, "研究分期：{{phase}}")
    _add_line(doc, "研究设计：{{design}}")
    _add_line(doc, "主要终点：{{primary_endpoint}}")
    _add_line(doc, "次要终点：{{secondary_endpoints}}")
    _add_line(doc, "计划样本量：{{sample_size}}")
    _add_line(doc, "研究周期：{{study_duration}}")
    _add_heading(doc, "2. 入排标准摘要", 1)
    _add_line(doc, "入选标准：{{inclusion_criteria}}")
    _add_line(doc, "排除标准：{{exclusion_criteria}}")
    _add_heading(doc, "3. 安全性监测", 1)
    _add_line(doc, "{{safety_monitoring}}")
    _add_heading(doc, "4. 期中分析与停止规则", 1)
    _add_line(doc, "期中分析：{{interim_analysis}}")
    _add_line(doc, "停止规则：{{stopping_rules}}")
    _add_heading(doc, "5. DMC 组织与运作", 1)
    _add_line(doc, "DMC 主席：{{dmc_chair}}")
    _add_line(doc, "DMC 成员：{{dmc_members}}")
    _add_line(doc, "会议频率：{{meeting_frequency}}")
    _add_heading(doc, "6. 保密条款", 1)
    _add_line(doc, "{{confidentiality}}")
    doc.save(path)
    return path


def _build_dmp() -> Path:
    path = SEED_DIR / "crogo_dmp.docx"
    doc = Document()
    _add_heading(doc, "数据管理计划 (DMP)", 0)
    _add_line(doc, "方案编号：{{protocol_id}}")
    _add_line(doc, "研究标题：{{study_title}}")
    _add_line(doc, "申办方：{{sponsor}}")
    _add_line(doc, "适应症：{{indication}}")
    _add_heading(doc, "1. 试验概述", 1)
    _add_line(doc, "研究分期：{{phase}}")
    _add_line(doc, "研究设计：{{design}}")
    _add_line(doc, "主要终点：{{primary_endpoint}}")
    _add_line(doc, "计划样本量：{{sample_size}}")
    _add_line(doc, "研究周期：{{study_duration}}")
    _add_heading(doc, "2. 数据采集与管理", 1)
    _add_line(doc, "入选标准：{{inclusion_criteria}}")
    _add_line(doc, "排除标准：{{exclusion_criteria}}")
    _add_line(doc, "安全性数据监测：{{safety_monitoring}}")
    _add_heading(doc, "3. 数据质量控制", 1)
    _add_line(doc, "期中分析相关数据要求：{{interim_analysis}}")
    _add_line(doc, "试验终止相关规则：{{stopping_rules}}")
    _add_heading(doc, "4. 数据库锁定与归档", 1)
    _add_line(doc, "数据管理由申办方 {{sponsor}} 负责，遵循方案 {{protocol_id}} 规定的数据标准与流程。")
    _add_heading(doc, "5. 保密与合规", 1)
    _add_line(doc, "{{confidentiality}}")
    doc.save(path)
    return path


def _build_sap() -> Path:
    path = SEED_DIR / "crogo_sap.docx"
    doc = Document()
    _add_heading(doc, "统计分析计划 (SAP)", 0)
    _add_line(doc, "方案编号：{{protocol_id}}")
    _add_line(doc, "研究标题：{{study_title}}")
    _add_line(doc, "申办方：{{sponsor}}")
    _add_line(doc, "适应症：{{indication}}")
    _add_heading(doc, "1. 研究背景与目的", 1)
    _add_line(doc, "研究分期：{{phase}}")
    _add_line(doc, "研究设计：{{design}}")
    _add_heading(doc, "2. 终点与估计目标", 1)
    _add_line(doc, "主要终点：{{primary_endpoint}}")
    _add_line(doc, "次要终点：{{secondary_endpoints}}")
    _add_heading(doc, "3. 样本量与把握度", 1)
    _add_line(doc, "计划样本量：{{sample_size}}")
    _add_line(doc, "研究周期：{{study_duration}}")
    _add_heading(doc, "4. 分析人群与入排标准", 1)
    _add_line(doc, "入选标准：{{inclusion_criteria}}")
    _add_line(doc, "排除标准：{{exclusion_criteria}}")
    _add_heading(doc, "5. 统计分析方法", 1)
    _add_line(doc, "主要终点将采用与方案一致的分析策略；安全性分析涵盖 {{safety_monitoring}}。")
    _add_heading(doc, "6. 期中分析与多重性", 1)
    _add_line(doc, "期中分析：{{interim_analysis}}")
    _add_line(doc, "停止规则：{{stopping_rules}}")
    _add_heading(doc, "7. 独立数据监查", 1)
    _add_line(doc, "DMC 主席：{{dmc_chair}}")
    _add_line(doc, "DMC 成员：{{dmc_members}}")
    _add_line(doc, "会议频率：{{meeting_frequency}}")
    doc.save(path)
    return path


TEMPLATE_SPECS = [
    ("DMC", "DMC Charter 标准模板", _build_dmc),
    ("DMP", "DMP 标准模板", _build_dmp),
    ("SAP", "SAP 标准模板", _build_sap),
]


def _default_mappings(placeholders: list[dict]) -> dict[str, str]:
    mappings: dict[str, str] = {}
    for ph in placeholders:
        name = ph["name"]
        if name in VALID_KEYS:
            mappings[name] = name
    return mappings


def _existing_names() -> set[str]:
    return {row["name"] for row in list_templates()}


def seed(*, force: bool = False) -> None:
    SEED_DIR.mkdir(parents=True, exist_ok=True)
    init_db()
    existing_names = _existing_names()

    for type_, name, builder in TEMPLATE_SPECS:
        if name in existing_names and not force:
            print(f"  skip {name} (already in database)")
            continue

        path = builder()
        sections, placeholders = parse_docx(path)
        desc = f"已解析 {len(sections)} 个章节、{len(placeholders)} 个占位符"
        tpl_id = create_template(type_, name, desc, str(path.resolve()), sections, placeholders)
        mappings = _default_mappings(placeholders)
        update_template_mappings(tpl_id, mappings)
        print(f"  + {type_} id={tpl_id} → {path.name} ({len(placeholders)} placeholders, mapped {len(mappings)})")

    # Copy seed files into templates dir for serving consistency
    templates_dir = settings.templates_dir
    templates_dir.mkdir(parents=True, exist_ok=True)
    for f in SEED_DIR.glob("crogo_*.docx"):
        dest = templates_dir / f.name
        if not dest.exists() or force:
            dest.write_bytes(f.read_bytes())


if __name__ == "__main__":
    force = "--force" in sys.argv
    print("Seeding Crogo templates...")
    seed(force=force)
    print("Done.")
