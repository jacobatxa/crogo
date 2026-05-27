from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

from models.schemas import FieldValue

PLACEHOLDER_PATTERNS = [
    (re.compile(r"\{\{([a-zA-Z0-9_]+)\}\}"), lambda k: f"{{{{{k}}}}}"),
    (re.compile(r"<([a-zA-Z0-9_\u4e00-\u9fff]+)>"), lambda k: f"<{k}>"),
    (re.compile(r"【([a-zA-Z0-9_\u4e00-\u9fff]+)】"), lambda k: f"【{k}】"),
    (re.compile(r"\[请[^\]]{0,60}填[^\]]*\]"), lambda k: k),
]


def _build_replacement_tokens(
    value_map: dict[str, str], mappings: dict[str, str]
) -> list[tuple[str, str]]:
    """Return (token, value) pairs sorted by token length descending."""
    pairs: list[tuple[str, str]] = []
    seen: set[str] = set()

    for ph_name, field_key in mappings.items():
        val = value_map.get(field_key, "") or value_map.get(ph_name, "")
        if not val:
            continue
        for pat, fmt in PLACEHOLDER_PATTERNS:
            token = fmt(ph_name)
            if token not in seen:
                seen.add(token)
                pairs.append((token, val))

    for key, val in value_map.items():
        if not val:
            continue
        for pat, fmt in PLACEHOLDER_PATTERNS:
            token = fmt(key)
            if token not in seen:
                seen.add(token)
                pairs.append((token, val))

    pairs.sort(key=lambda x: len(x[0]), reverse=True)
    return pairs


def _replace_in_text(text: str, replacements: list[tuple[str, str]]) -> str:
    result = text
    for token, val in replacements:
        if token and token in result:
            result = result.replace(token, val)
    return result


def _apply_paragraph_replace(paragraph: Paragraph, replacements: list[tuple[str, str]]) -> bool:
    full = "".join(r.text for r in paragraph.runs)
    if not full.strip():
        return False
    new_text = _replace_in_text(full, replacements)
    if new_text == full:
        return False
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = new_text
    else:
        paragraph.add_run(new_text)
    return True


def generate_docx(
    template_path: Path,
    output_path: Path,
    fields: list[FieldValue],
    mappings: dict[str, str],
) -> None:
    value_map = {f.key: f.value for f in fields if f.value}
    for ph, fk in mappings.items():
        if fk in value_map and value_map[fk]:
            value_map[ph] = value_map[fk]

    replacements = _build_replacement_tokens(value_map, mappings)
    doc = Document(template_path)

    for para in doc.paragraphs:
        _apply_paragraph_replace(para, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _apply_paragraph_replace(para, replacements)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
