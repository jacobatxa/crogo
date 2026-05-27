from __future__ import annotations

import re
from pathlib import Path

from docx import Document
PLACEHOLDER_PATTERNS = [
    re.compile(r"\{\{([a-zA-Z0-9_]+)\}\}"),
    re.compile(r"<([a-zA-Z0-9_\u4e00-\u9fff]+)>"),
    re.compile(r"【([a-zA-Z0-9_\u4e00-\u9fff]+)】"),
    re.compile(r"\[请[^\]]{0,60}填[^\]]*\]"),
]


def _extract_placeholders_from_text(text: str) -> list[str]:
    found: list[str] = []
    for pat in PLACEHOLDER_PATTERNS:
        for m in pat.finditer(text):
            if pat.pattern.startswith(r"\[请"):
                name = m.group(0)
            else:
                name = m.group(1)
            if name not in found:
                found.append(name)
    return found


def parse_docx(path: Path) -> tuple[list[dict], list[dict]]:
    doc = Document(path)
    sections: list[dict] = []
    placeholders: list[dict] = []
    seen_ph: set[str] = set()

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name if para.style else ""
        level = 0
        if style_name.startswith("Heading"):
            try:
                level = int(style_name.replace("Heading", "").strip() or "1")
            except ValueError:
                level = 1

        if level > 0 or re.match(r"^\d+(\.\d+)*\s+\S", text):
            number = ""
            m = re.match(r"^(\d+(?:\.\d+)*)\s+", text)
            if m:
                number = m.group(1)
            sections.append(
                {
                    "level": level or 1,
                    "title": text,
                    "number": number,
                }
            )

        for name in _extract_placeholders_from_text(text):
            if name not in seen_ph:
                seen_ph.add(name)
                placeholders.append({"name": name, "context": text[:120]})

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    t = para.text.strip()
                    if not t:
                        continue
                    for name in _extract_placeholders_from_text(t):
                        if name not in seen_ph:
                            seen_ph.add(name)
                            placeholders.append({"name": name, "context": t[:120]})

    return sections, placeholders
