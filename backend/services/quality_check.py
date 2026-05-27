from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx import Document

RESIDUAL_PATTERNS = [
    re.compile(r"\{\{[a-zA-Z0-9_]+\}\}"),
    re.compile(r"<[a-zA-Z0-9_\u4e00-\u9fff]+>"),
    re.compile(r"【[a-zA-Z0-9_\u4e00-\u9fff]+】"),
    re.compile(r"\[请[^\]]{0,40}填[^\]]*\]"),
]


def check_generated_docx(
    path: Path,
    *,
    expected_placeholders: list[dict] | None = None,
    mappings: dict[str, str] | None = None,
) -> dict[str, Any]:
    doc = Document(path)
    texts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            texts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        texts.append(para.text)

    full = "\n".join(texts)
    issues: list[dict[str, str]] = []

    for pat in RESIDUAL_PATTERNS:
        for m in pat.finditer(full):
            snippet = m.group(0)
            if not any(i["content"] == snippet for i in issues):
                issues.append(
                    {
                        "severity": "high",
                        "type": "unfilled_placeholder",
                        "content": snippet,
                        "suggestion": "检查映射或方案字段是否为空",
                    }
                )

    total_ph = len(expected_placeholders or [])
    filled_estimate = max(0, total_ph - len(issues)) if total_ph else 1
    fill_rate = filled_estimate / total_ph if total_ph else 1.0
    requires_review = len(issues)
    grade = _grade(fill_rate, requires_review)

    return {
        "grade": grade,
        "fill_rate": round(fill_rate, 3),
        "total_placeholders": total_ph,
        "mapped_count": len(mappings or {}),
        "requires_review": requires_review,
        "issues": issues[:20],
    }


def _grade(fill_rate: float, requires_review: int) -> str:
    if fill_rate >= 0.95 and requires_review == 0:
        return "S"
    if fill_rate >= 0.8 and requires_review <= 10:
        return "A"
    if fill_rate >= 0.6:
        return "B"
    return "C"


def format_quality_message(report: dict[str, Any]) -> str:
    return (
        f"等级 {report.get('grade', '?')} · "
        f"填充率 {int((report.get('fill_rate') or 0) * 100)}% · "
        f"需确认 {report.get('requires_review', 0)} 项"
    )
